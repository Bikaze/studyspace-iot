"""
Generate the StudySpace IoT project report as a .docx file.
Run from the repo root:  python3 generate_report.py
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = Path(__file__).parent
IMAGES  = BASE / "images"
SCREENS = BASE / "analysis" / "screenshots"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Styles helpers ────────────────────────────────────────────────────────────

def para(text="", bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
         color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after  = Pt(6)
    return h

def add_image(path, width=6.0, caption=None):
    if Path(path).exists():
        doc.add_picture(str(path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
            for run in cp.runs:
                run.italic = True
                run.font.size = Pt(9)

def add_code(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EFF6FF")
    p._p.get_or_add_pPr().append(shd)
    return p

def divider():
    p = doc.add_paragraph("─" * 80)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# University logo
logo_path = IMAGES / "ur_logo.png"
if logo_path.exists():
    doc.add_picture(str(logo_path), width=Inches(1.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
para("UNIVERSITY OF RWANDA", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER,
     color=(0x1e, 0x3a, 0x8a))
para("College of Science and Technology", bold=False, size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x37, 0x51, 0x9a))
para("Department of Computer Science", bold=False, size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x37, 0x51, 0x9a))
doc.add_paragraph()
para("Course: Ubiquitous and Pervasive Computing", italic=True, size=11,
     align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

para("StudySpace IoT", bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER,
     color=(0x0f, 0x17, 0x2a), space_before=20)
para("Smart Study Room Monitoring System", bold=False, size=14,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x47, 0x55, 0x69))
doc.add_paragraph()
para("Project Report", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()

# Group members table
para("Group 4 — Members", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=20)
doc.add_paragraph()

members_table = doc.add_table(rows=13, cols=2)
members_table.style = "Table Grid"
members_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = members_table.rows[0].cells
hdr[0].text = "Registration Number"
hdr[1].text = "Full Name"
for cell in hdr:
    for para_ in cell.paragraphs:
        for run in para_.runs:
            run.bold = True
            run.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1E3A8A")
    cell._tc.get_or_add_tcPr().append(shd)
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

for i in range(1, 13):
    for cell in members_table.rows[i].cells:
        cell.text = ""
        cell.paragraphs[0].paragraph_format.space_after = Pt(8)

members_table.columns[0].width = Inches(2.5)
members_table.columns[1].width = Inches(3.5)

doc.add_paragraph()
para(f"Academic Year: 2025 / 2026", size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
     italic=True, color=(0x64, 0x74, 0x8b))
para(f"Date: May 2026", size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
     italic=True, color=(0x64, 0x74, 0x8b))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

heading("1. Introduction", level=1)
para(
    "StudySpace IoT is a full-stack IoT system designed to monitor and improve "
    "the study environment in university classrooms and libraries. Five environmental "
    "sensors mounted on an ESP32 microcontroller continuously measure temperature, "
    "humidity, ambient sound, light intensity, and occupancy movement. Every five "
    "seconds these readings are transmitted over WiFi to a FastAPI backend, stored "
    "in PostgreSQL, and made available through a React dashboard for students and "
    "facility managers.",
    size=11
)
para(
    "The system goes beyond raw data collection. Each reading is automatically "
    "classified into one of eight comfort categories (comfortable, humid, noisy, "
    "crowded, warm, dim, moderate, poor) using a rule engine running on the backend. "
    "A composite comfort score (0–100) is computed from three physiologically-grounded "
    "components: thermal comfort, acoustic comfort, and visual comfort. Machine learning "
    "models trained on this labelled dataset then power the Insights page — providing "
    "real-time predictions and feature importance explanations.",
    size=11
)

para("System Architecture", bold=True, size=12, space_before=10)
para(
    "The system follows a layered architecture: the firmware layer collects and "
    "transmits sensor data; the backend layer transforms, stores, and classifies "
    "readings; the frontend layer visualises and presents the data; and the analysis "
    "layer trains machine learning models offline and exports them back to the backend.",
    size=11
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — FIRMWARE
# ══════════════════════════════════════════════════════════════════════════════

heading("2. Firmware — ESP32 Sensor Node", level=1)
para(
    "The firmware is written in C++ for the Arduino framework and runs on an "
    "ESP32 development board. It coordinates four sensors, handles WiFi connectivity, "
    "synchronises time with NTP, and sends sensor payloads to the backend every "
    "five seconds.",
    size=11
)

# ESP32 image
add_image(IMAGES / "esp32.jpeg", width=3.5, caption="Figure 1 — ESP32 Development Board")

heading("2.1 Sensors", level=2)
para(
    "Each sensor is chosen for its suitability in an indoor study environment. "
    "The table below lists the sensors, their measurement targets, and the GPIO "
    "pins used.",
    size=11
)

sensor_tbl = doc.add_table(rows=6, cols=3)
sensor_tbl.style = "Table Grid"
headers = ["Sensor", "Measures", "GPIO Pins"]
for i, h in enumerate(headers):
    cell = sensor_tbl.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "DBEAFE")
    cell._tc.get_or_add_tcPr().append(shd)

rows_data = [
    ("DHT22",    "Temperature (°C) and Relative Humidity (%)",  "GPIO 4 (single-wire, 10 kΩ pull-up to 3.3 V)"),
    ("HC-SR501", "Passive infrared motion (PIR) — counting movement events", "GPIO 27 (rising-edge interrupt)"),
    ("GL5528 LDR", "Ambient light via voltage divider",         "GPIO 34 (ADC1, input-only)"),
    ("INMP441",  "Sound pressure via I2S digital microphone",   "GPIO 14 (SCK), 15 (WS), 32 (SD)"),
]
for i, (s, m, g) in enumerate(rows_data, start=1):
    sensor_tbl.rows[i].cells[0].text = s
    sensor_tbl.rows[i].cells[1].text = m
    sensor_tbl.rows[i].cells[2].text = g
    for cell in sensor_tbl.rows[i].cells:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)

doc.add_paragraph()

# Sensor images
heading("2.2 Sensor Hardware", level=2)

sensor_imgs = [
    (IMAGES / "temp_humidity_sensor.png",  "Figure 2 — DHT22 Temperature & Humidity Sensor"),
    (IMAGES / "sound_sensor.png",          "Figure 3 — INMP441 I2S Microphone"),
    (IMAGES / "light_sensor.png",          "Figure 4 — GL5528 LDR Light Sensor"),
    (IMAGES / "motion_sensor.png",         "Figure 5 — HC-SR501 PIR Motion Sensor"),
]
for img_path, caption in sensor_imgs:
    add_image(img_path, width=3.0, caption=caption)

heading("2.3 Wiring", level=2)
para(
    "All sensors operate at 3.3 V logic, which matches the ESP32's GPIO voltage. "
    "The key wiring points are:", size=11
)
wiring_items = [
    "DHT22: DATA → GPIO 4, with a 10 kΩ pull-up resistor between DATA and 3.3 V. "
     "Without the pull-up the single-wire protocol fails.",
    "HC-SR501: OUTPUT → GPIO 27 configured as a rising-edge interrupt so every "
     "detected motion burst increments a counter in IRAM without blocking the main loop.",
    "GL5528 LDR: Wired as a voltage divider with a 10 kΩ fixed resistor between 3.3 V "
     "and GND. The ADC reads the mid-point. Higher light → lower LDR resistance → "
     "higher voltage → higher ADC count.",
    "INMP441: SCK → GPIO 14, WS → GPIO 15, SD → GPIO 32. The L/R pin is tied to GND "
     "so the mic always outputs on the left I2S channel.",
]
for item in wiring_items:
    p = doc.add_paragraph(item, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)

heading("2.4 Configuration (config.h)", level=2)
para(
    "All user-configurable constants live in a single header file so the firmware "
    "can be re-flashed for any room without touching the core logic.", size=11
)
add_code(
    "#define WIFI_SSID        \"your_wifi_name\"\n"
    "#define WIFI_PASSWORD    \"your_wifi_password\"\n"
    "#define BACKEND_URL      \"http://192.168.x.x:8000/api/ingest\"\n"
    "#define ROOM_ID          \"your_room_slug\"\n"
    "#define SEND_INTERVAL    5000     // ms between readings\n"
    "#define DHT_PIN          4\n"
    "#define PIR_PIN          27\n"
    "#define LDR_PIN          34\n"
    "#define I2S_SCK_PIN      14\n"
    "#define I2S_WS_PIN       15\n"
    "#define I2S_SD_PIN       32\n"
    "#define I2S_SAMPLE_COUNT 1024    // ~23 ms of audio at 44100 Hz\n"
    "#define CAT_OFFSET_SEC   7200    // UTC+2, no DST"
)

heading("2.5 Main Loop", level=2)
para(
    "The main loop uses non-blocking timing so the PIR interrupt counter is never "
    "frozen by a long delay. Every 5 seconds it:", size=11
)
loop_steps = [
    "Reads all sensors via readSensors()",
    "Builds a JSON payload manually (no ArduinoJson dependency needed)",
    "Checks WiFi is still connected before attempting the HTTP call",
    "HTTP POSTs the payload to /api/ingest and logs the response code",
    "Calls http.end() to release the TCP socket — skipping this exhausts the "
     "ESP32 socket table after ~4 hours",
]
for step in loop_steps:
    p = doc.add_paragraph(step, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(11)

heading("2.6 Key Code: Reading the INMP441 Microphone", level=2)
para(
    "The microphone uses the I2S digital protocol. The firmware reads 1024 samples "
    "into a buffer, then computes the Root Mean Square (RMS) amplitude. The INMP441 "
    "left-justifies 24-bit audio in a 32-bit I2S frame, so each sample is "
    "right-shifted by 8 before squaring to remove the zero-padding.", size=11
)
add_code(
    "int32_t samples[I2S_SAMPLE_COUNT];\n"
    "size_t bytesRead = 0;\n"
    "i2s_read(I2S_NUM_0, samples, sizeof(samples), &bytesRead, portMAX_DELAY);\n\n"
    "int64_t sumSquares = 0;\n"
    "for (int i = 0; i < samplesRead; i++) {\n"
    "    int32_t sample = samples[i] >> 8;  // remove 8-bit zero padding\n"
    "    sumSquares += (int64_t)sample * sample;\n"
    "}\n"
    "data.soundRms = (int)sqrt((double)sumSquares / samplesRead);"
)

heading("2.7 NTP Time Synchronisation", level=2)
para(
    "The ESP32 has no real-time clock battery. On first boot it contacts "
    "pool.ntp.org to sync its internal clock to Central Africa Time (UTC+2). "
    "Every payload timestamp is formatted as ISO 8601 with an explicit +02:00 "
    "offset so the backend's Pydantic parser interprets it correctly and "
    "PostgreSQL stores it normalised to UTC.", size=11
)
add_code('strftime(buf, sizeof(buf), "%Y-%m-%dT%H:%M:%S+02:00", &timeinfo);')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — BACKEND
# ══════════════════════════════════════════════════════════════════════════════

heading("3. Backend — FastAPI & PostgreSQL", level=1)
para(
    "The backend is a Python FastAPI application backed by a PostgreSQL database. "
    "It receives raw sensor payloads from the ESP32, applies physical unit conversions, "
    "computes a comfort score, classifies the reading, detects anomalies, and exposes "
    "a REST API consumed by the React frontend.", size=11
)

heading("3.1 Tech Stack", level=2)
stack = [
    ("FastAPI", "Async Python web framework — automatic OpenAPI docs at /docs"),
    ("SQLAlchemy (async)", "ORM for database access with asyncpg driver"),
    ("PostgreSQL", "Time-series sensor readings storage"),
    ("Pydantic v2", "Request/response validation and serialisation"),
    ("joblib + scikit-learn", "Load trained ML models at startup for /predict endpoint"),
]
tbl = doc.add_table(rows=len(stack)+1, cols=2)
tbl.style = "Table Grid"
for i, h in enumerate(["Component", "Role"]):
    tbl.rows[0].cells[i].text = h
    for run in tbl.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True; run.font.size = Pt(10)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "DBEAFE")
    tbl.rows[0].cells[i]._tc.get_or_add_tcPr().append(shd)
for i, (comp, role) in enumerate(stack, 1):
    tbl.rows[i].cells[0].text = comp
    tbl.rows[i].cells[1].text = role
    for cell in tbl.rows[i].cells:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
doc.add_paragraph()

heading("3.2 Sensor Transforms (transforms.py)", level=2)
para(
    "This module converts raw ESP32 values into meaningful physical units. "
    "It has no database imports — every function is independently testable.", size=11
)

transforms = [
    ("adc_to_lux(adc_value)",
     "Converts the 12-bit LDR ADC count to lux using the GL5528 inverse power-law: "
     "lux ≈ 500 / (R_kΩ ^ 0.7). First recovers the voltage from the ADC count, "
     "then back-calculates the LDR resistance from the voltage divider equation, "
     "then applies the empirical curve."),
    ("rms_to_db(rms_value)",
     "Converts raw INMP441 RMS amplitude to dB SPL. Anchored so that 420,426 RMS "
     "reads as 94 dB SPL (INMP441 datasheet sensitivity at 1 kHz, 1 Pa). "
     "Formula: dB_SPL = 20 × log10(rms / 420426) + 94"),
    ("compute_movements_per_min(motion_count)",
     "Scales the 5-second PIR interrupt count to per-minute rate: count × 12 "
     "(because 60 s / 5 s = 12 windows per minute)."),
    ("apparent_temperature(T, RH)",
     "Australian Bureau of Meteorology formula: AT = T + 0.33 × e − 4.0 "
     "where e is partial vapour pressure. Combines temperature and humidity into "
     "a single physiological heat perception value."),
    ("compute_comfort_score(...)",
     "Three-component composite score (0–100): Thermal (40 pts) from apparent "
     "temperature, Acoustic (35 pts) from sound dB with crowding amplifier, "
     "Visual (25 pts) from lux. Based on ASHRAE 55-2023 and WHO Noise Guidelines."),
    ("classify_reading(...)",
     "Applies priority-ordered rules to assign one of eight labels: poor, warm, "
     "humid, noisy, dim, crowded, moderate, comfortable."),
    ("run_all_transforms(payload, thresholds)",
     "Single entry point called by the ingest route — runs all conversions and "
     "returns a dict ready to merge into the database row."),
]

for fn, desc in transforms:
    p = doc.add_paragraph()
    run = p.add_run(fn)
    run.bold = True; run.font.name = "Courier New"; run.font.size = Pt(10)
    run2 = p.add_run(f" — {desc}")
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

heading("3.3 API Endpoints", level=2)

endpoints = [
    ("POST /api/ingest",
     "Main endpoint consumed by the ESP32. Validates the payload with Pydantic, "
     "confirms the room_id exists, runs all transforms, writes the reading to "
     "sensor_readings, then runs anomaly detection and writes any flags to the "
     "anomalies table. Returns HTTP 201 on success."),
    ("GET /api/rooms",
     "Lists all registered rooms ordered by creation date."),
    ("POST /api/rooms",
     "Creates a new room. The room name is automatically slugified "
     "(e.g. 'Library Floor 2' → 'library_floor_2') for use as the room_id."),
    ("DELETE /api/rooms/{room_id}",
     "Removes a room and all associated readings (cascade)."),
    ("GET /api/rooms/{room_id}/readings",
     "Returns paginated raw readings for a room (limit/offset query params)."),
    ("GET /api/rooms/{room_id}/latest",
     "Returns the single most recent reading — used by the dashboard live tile."),
    ("GET /api/rooms/{room_id}/summary",
     "Returns avg/min/max for all six metrics over the last 24 hours."),
    ("GET /api/rooms/{room_id}/label-distribution",
     "Returns label counts for the last 24 hours — used by pie/bar charts."),
    ("GET /api/rooms/{room_id}/correlation",
     "Computes and returns the 5×5 Pearson correlation matrix for the last N readings."),
    ("GET /api/rooms/{room_id}/predict",
     "Runs the latest reading through the trained Random Forest classifier and "
     "returns the predicted label, confidence, and feature importances."),
    ("GET /api/anomalies",
     "Returns recent anomaly events with metric name, value, and plain-English reason."),
    ("GET /api/thresholds",
     "Returns the current comfort thresholds used in scoring and classification."),
    ("PUT /api/thresholds",
     "Updates thresholds — changes take effect immediately for all future readings."),
]

ep_tbl = doc.add_table(rows=len(endpoints)+1, cols=2)
ep_tbl.style = "Table Grid"
for i, h in enumerate(["Endpoint", "Description"]):
    ep_tbl.rows[0].cells[i].text = h
    for run in ep_tbl.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True; run.font.size = Pt(10)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "DBEAFE")
    ep_tbl.rows[0].cells[i]._tc.get_or_add_tcPr().append(shd)
for i, (ep, desc) in enumerate(endpoints, 1):
    ep_tbl.rows[i].cells[0].text = ep
    ep_tbl.rows[i].cells[1].text = desc
    ep_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.name = "Courier New"
    ep_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    for run in ep_tbl.rows[i].cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

doc.add_paragraph()
heading("3.4 Anomaly Detection", level=2)
para(
    "After every successful ingest, the backend checks whether any metric exceeds "
    "wider 'anomaly bounds' — deliberately set beyond the comfort thresholds to "
    "catch physically unusual events rather than minor discomfort:", size=11
)
anomalies_items = [
    "Apparent temperature > temp_max + 5 °C or < temp_min − 5 °C → possible HVAC failure",
    "Humidity > 78% → condensation/mould risk | < 28% → excessively dry air",
    "Sound > sound_max + 18 dB → acoustic spike event",
    "Light < 100 lux → lamp failure | > 900 lux → direct sunlight / fixture fault",
    "Movement > motion_max × 3 → unusual occupancy event",
]
for item in anomalies_items:
    p = doc.add_paragraph(item, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 3 — FRONTEND
# ══════════════════════════════════════════════════════════════════════════════

heading("4. Frontend — React Dashboard", level=1)
para(
    "The dashboard is built with React, TypeScript, and Tailwind CSS. It runs "
    "entirely in the browser and communicates with the backend through the REST API. "
    "Five pages cover all monitoring and management needs.", size=11
)

pages = [
    ("Rooms (Home)",
     "Lists all registered rooms with their current comfort score and label. "
     "Colour-coded badges show the label at a glance. Clicking a room opens the Room Detail page.",
     IMAGES / "frontend_rooms_page.png",
     "Figure 6 — Rooms overview page"),
    ("Room Detail",
     "Shows live sensor readings for a selected room: temperature, humidity, sound, "
     "light, and motion tiles update from the /latest endpoint. Includes a 24-hour "
     "summary table with avg/min/max for every metric and a label distribution chart.",
     IMAGES / "frontend_room_details_page.png",
     "Figure 7 — Room detail page"),
    ("Metric Detail",
     "Drill-down page for a single metric. Shows a time-series chart of the last "
     "N readings and the 24-hour summary statistics.",
     IMAGES / "frontend_temperature_metric_example.png",
     "Figure 8 — Temperature metric detail page"),
    ("Insights",
     "ML-powered page. Fetches predictions from /predict and displays the predicted "
     "label, confidence percentage, and a horizontal bar chart of feature importances "
     "showing which sensor drove the prediction.",
     IMAGES / "frontend_insights_page.png",
     "Figure 9 — Insights page with ML predictions"),
    ("Anomalies",
     "Lists all anomaly events in reverse chronological order. Each card shows the "
     "metric name, the measured value, and a plain-English reason generated by the "
     "backend's anomaly detection logic.",
     IMAGES / "frontend_anomalies_page.png",
     "Figure 10 — Anomalies page"),
    ("Settings",
     "Form to update the comfort thresholds (temp range, humidity, sound, light, "
     "motion limits). Changes take effect immediately via PUT /api/thresholds.",
     IMAGES / "frontend_settings_page.png",
     "Figure 11 — Settings page"),
]

for title, desc, img, caption in pages:
    heading(f"4.{pages.index((title, desc, img, caption))+1} {title}", level=2)
    para(desc, size=11)
    add_image(img, width=5.5, caption=caption)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 4 — ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

heading("5. Data Analysis & Machine Learning", level=1)
para(
    "The analysis is performed in a Jupyter Notebook (analysis/studyspace_analysis.ipynb). "
    "It runs once after sufficient data is collected and saves two model files to "
    "the backend — comfort_classifier.pkl and feature_scaler.pkl — which the "
    "Insights page then uses for live predictions.", size=11
)

heading("5.1 Pipeline Overview", level=2)
add_code(
    "PostgreSQL  →  Clean  →  Normalise  →  Visualise  →  Correlate\n"
    "           →  Detect Outliers  →  Train Classifiers  →  Train Regressors\n"
    "           →  Save .pkl files  →  Backend serves predictions"
)

heading("5.2 Data Loading & Cleaning", level=2)
para(
    "Data is loaded from PostgreSQL using pandas read_sql with the psycopg2 driver "
    "(the asyncpg driver used by FastAPI is incompatible with synchronous pandas calls). "
    "After loading, rows with missing sensor values are dropped and time features "
    "(hour, weekday, is_weekday) are added. Labels are encoded as integers for ML "
    "in alphabetical order: comfortable=0, crowded=1, dim=2, humid=3, "
    "moderate=4, noisy=5, poor=6, warm=7.", size=11
)

heading("5.3 Normalisation", level=2)
para(
    "All five sensor features are scaled to [0, 1] using Min-Max normalisation:", size=11
)
add_code("x_scaled = (x − x_min) / (x_max − x_min)")
para(
    "The scaler is fitted once on the training data and saved alongside the model. "
    "At prediction time, the same scaler is applied to new readings — this ensures "
    "the model always sees the same scale it was trained on.", size=11
)

heading("5.4 Time-Series Visualisation", level=2)
para("A sanity check — plots the last 500 readings to confirm sensors behave as expected before training.", size=11)
add_image(SCREENS / "time_series_visualization.png", width=5.5,
          caption="Figure 12 — Time-series sensor readings (last 500 records)")

heading("5.5 Correlation Analysis", level=2)
para(
    "A Pearson correlation heatmap identifies linear relationships between sensors. "
    "Correlation coefficient r ranges from −1 (perfect inverse) to +1 (perfect positive). "
    "Values near 0 mean the sensors are independent. This informs feature selection "
    "and helps explain why certain labels co-occur.", size=11
)
add_image(SCREENS / "correlation_analysis.png", width=4.5,
          caption="Figure 13 — Pearson correlation matrix")

heading("5.6 Outlier Detection (IQR Method)", level=2)
para(
    "Outliers are flagged using Tukey's IQR fences: any value below Q1 − 1.5×IQR "
    "or above Q3 + 1.5×IQR is an outlier. Approximately 11% of readings are flagged. "
    "The clean dataset (outliers removed) is used for regression modelling.", size=11
)
add_image(SCREENS / "outlier_detection.png", width=5.5,
          caption="Figure 14 — Box plots showing outlier distribution per metric")

heading("5.7 Classification Models", level=2)
para(
    "Three classifiers are trained on the labelled dataset to predict one of "
    "eight comfort categories from the five sensor features.", size=11
)

clf_data = [
    ("Random Forest",
     "An ensemble of 100 decision trees. Each tree votes for a label and the "
     "majority wins. Provides feature importances (what fraction of splits "
     "each sensor contributed). Selected as the production model.",
     "99.81%"),
    ("Logistic Regression",
     "A linear model that multiplies each feature by a learned weight, sums "
     "them, then applies the softmax function to produce a probability per class. "
     "Fast and interpretable but assumes linear decision boundaries.",
     "93.90%"),
    ("Decision Tree",
     "A single tree of up to 6 levels of yes/no questions. Fully transparent — "
     "the exact rules can be visualised. Good performance with no hyperparameter tuning.",
     "99.12%"),
]

for name, desc, acc in clf_data:
    p = doc.add_paragraph()
    run = p.add_run(f"{name} ({acc} accuracy) — ")
    run.bold = True; run.font.size = Pt(11)
    run2 = p.add_run(desc)
    run2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

add_image(SCREENS / "random_forest_training_output.png", width=5.5,
          caption="Figure 15 — Random Forest classification report")
add_image(SCREENS / "model_comparison.png", width=5.5,
          caption="Figure 16 — Model accuracy and per-class F1 comparison")

heading("5.8 Confusion Matrix", level=2)
para(
    "The confusion matrix shows where the best model (Random Forest) makes mistakes. "
    "Rows are actual labels, columns are predicted labels. The diagonal is correct "
    "predictions. Off-diagonal cells are errors. With 99.81% accuracy, almost all "
    "counts sit on the diagonal.", size=11
)
add_image(SCREENS / "confusion_matrix.png", width=4.5,
          caption="Figure 17 — Random Forest confusion matrix")

heading("5.9 Regression Models", level=2)
para(
    "Two regression models predict comfort_score (0–100) as a continuous number:", size=11
)
reg_data = [
    ("Linear Regression",
     "Fits a weighted sum of the five features. All coefficients are negative — "
     "every sensor going up hurts comfort. Sound has the largest impact (−114.4). "
     "R²=0.81, RMSE=7.4 points."),
    ("Random Forest Regressor",
     "100 trees each predict a score; the average is the final output. Captures "
     "non-linear relationships. R²=0.978, RMSE=2.5 points — off by only 2.5 "
     "comfort points on average across the 0–100 scale."),
]
for name, desc in reg_data:
    p = doc.add_paragraph()
    run = p.add_run(f"{name} — ")
    run.bold = True; run.font.size = Pt(11)
    run2 = p.add_run(desc)
    run2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

add_image(SCREENS / "predicting_comfort_score.png", width=5.5,
          caption="Figure 18 — Actual vs predicted comfort score (Linear vs Random Forest Regressor)")

heading("5.10 Saving Models", level=2)
para(
    "The trained Random Forest classifier and the fitted MinMaxScaler are serialised "
    "with joblib and saved to backend/models/. Both files are required: the model "
    "for predictions, the scaler to apply the same normalisation to new readings "
    "that was used during training.", size=11
)
add_code(
    "joblib.dump(rf_clf,  '../backend/models/comfort_classifier.pkl')\n"
    "joblib.dump(scaler,  '../backend/models/feature_scaler.pkl')"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

heading("6. Conclusion", level=1)
para(
    "StudySpace IoT demonstrates a complete IoT pipeline — from hardware sensor "
    "nodes through a cloud-connected backend to a machine-learning-powered dashboard. "
    "The system successfully classifies study room conditions in real time with "
    "99.81% accuracy using a Random Forest model trained on 207,000 labelled readings, "
    "and predicts the composite comfort score to within 2.5 points on a 100-point scale.",
    size=11
)
para(
    "The rule-based classification engine and the ML model agree on 99.87% of readings, "
    "validating both approaches independently. The three-component comfort score "
    "(thermal 40 pts, acoustic 35 pts, visual 25 pts) is grounded in ASHRAE 55-2023, "
    "WHO Environmental Noise Guidelines, and EN 12464-1:2021, providing a scientifically "
    "defensible measure of study environment quality.",
    size=11
)
para(
    "Future work could include CO₂ sensors for air quality, multi-room comparative "
    "dashboards, push notification alerts, and student feedback integration to "
    "validate and refine the comfort model over time.",
    size=11
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = BASE / "StudySpace_IoT_Report.docx"
doc.save(str(out))
print(f"Report saved → {out}")
